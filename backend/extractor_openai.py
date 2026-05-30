import pdfplumber
import os
import re
import json
import logging
import tempfile
import requests
import time
from typing import Dict, Any, List, Optional
from dotenv import load_dotenv
from groq import Groq
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_classic.chains import RetrievalQA
from langchain_groq import ChatGroq
from langchain_huggingface import HuggingFaceEmbeddings

# Load GROQ_API_KEY from .env
load_dotenv()

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ============================================================
# SPECIAL KEYWORDS — detected via regex in full PDF text
# ============================================================
SPECIAL_KEYWORDS = {
    "Malicious Code Certificate (MCC)": [
        r"malicious[\s\-]*code[\s\-]*certificate", r"\bMCC\b"
    ],
    "Non Disclosure Declaration (NDD)": [
        r"non[\s\-]*disclosure[\s\-]*declaration", r"\bNDD\b"
    ],
    "MII Compliance (Make in India Compliance)": [
        r"MII[\s\-]*compliance", r"make[\s\-]*in[\s\-]*india[\s\-]*compliance"
    ],
    "Relaxation": [r"\brelaxation\b"],
    "International Organization for Standardization (ISO)": [
        r"international[\s\-]*organization[\s\-]*for[\s\-]*standardization",
        r"\bISO\b", r"ISO[\s\-]?\d+"
    ],
    "Bureau of Indian Standards (BIS)": [
        r"bureau[\s\-]*of[\s\-]*indian[\s\-]*standards?", r"\bBIS\b"
    ],
    "Mandate":        [r"\bmandate\b"],
    "NET WORTH":      [r"net[\s\-]*worth"],
    "Escalation":     [r"\bescalation\b"],
    "Matrix":         [r"\bmatrix\b"],
    "Exemption":      [r"\bexemption\b"],
    "meeting, meet":  [r"\bmeeting\b", r"\bmeet\b"],
    "Pre Bid Detail(s)": [r"pre[\s\-]*bid[\s\-]*detail(?:s)?"],
    "Judicial / Stamp Paper": [r"\bjudicial\b", r"stamp[\s\-]*paper"],
    "Notarised / Notarized": [r"\bnotarised\b", r"\bnotarized\b"],
    "Pre Bid Queries": [r"pre[\s\-]*bid[\s\-]*queries"],
    "Exempted": [r"\bexempted\b"],
    "Blacklisted / Debarred": [
        r"black[\s\-]*listing", r"black[\s\-]*listed", r"\bblacklisting\b", 
        r"\bblacklisted\b", r"\bdebarred\b"
    ],
    "Fraud / Criminal": [r"\bcriminal\b", r"\bfraudulent\b", r"\bfraud\b"],
    "Liquidation / Court Receivership": [r"\bliquidation\b", r"court[\s\-]*receivership"]
}

# Singleton embedding model
_embedding_model = None

def get_embedding_model():
    global _embedding_model
    if _embedding_model is None:
        logger.info("Loading embedding model (first time)...")
        _embedding_model = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2"
        )
    return _embedding_model


class CompleteTenderExtractor:
    """
    High-Accuracy GeM Extractor with human-like document identification.
    """

    def __init__(self):
        self.api_key = os.getenv("GROQ_API_KEY")
        if not self.api_key:
            raise ValueError("GROQ_API_KEY environment variable is not set.")
        self.client = Groq(api_key=self.api_key)

    def extract_fields(self, pdf_path: str) -> Dict[str, Any]:
        """Main entry point for bid document extraction."""
        relevant_pages = self._get_relevant_context(pdf_path)
        raw_full_context = "\n".join([p['raw_text'] for p in relevant_pages])
        full_context = "\n".join([p['text'] for p in relevant_pages])

        if not raw_full_context:
            return self._default_data()

        data = self._extract_basic_info(raw_full_context)
        logger.info(f"[CLASSIFY] Detected classification_level = '{data['classification_level']}'")

        if data.get("classification_level") == "Q1":
            logger.info("[CLASSIFY] Q1 tender -> returning immediately.")
            return data
            
        data["keyword_checks"] = self._check_keywords(raw_full_context)

        # Main AI extraction
        if self.client:
            try:
                logger.info(f"🤖 Analyzing main tender with Llama 3.3 70B...")
                ai_result = self._deep_ai_extraction(full_context, is_atc=False)
                data["documents_by_section"] = ai_result.get("documents_by_section", data["documents_by_section"])
            except Exception as e:
                logger.warning(f"⚠️ Main AI Pass failed: {e}")
                data = self._regex_fallback(full_context, data)

        # ATC processing
        atc_link = self._extract_atc_link(pdf_path)
        data["atc_document_link"] = atc_link
        if atc_link:
            atc_text = self._fetch_atc_text(atc_link)
            if atc_text:
                logger.info("📎 Executing AI pass for ATC document.")
                atc_kw = self._check_keywords(atc_text)
                for k, v in atc_kw.items():
                    if v:
                        data["keyword_checks"][k] = True
                
                if self.client:
                    try:
                        atc_result = self._deep_ai_extraction(atc_text, is_atc=True)
                        atc_docs = atc_result.get("documents_by_section", {})
                        if "atc" not in data["documents_by_section"]:
                            data["documents_by_section"]["atc"] = []
                        for docs in atc_docs.values():
                            for doc in docs:
                                if doc not in data["documents_by_section"]["atc"]:
                                    data["documents_by_section"]["atc"].append(doc)
                        total_atc_added = len(data["documents_by_section"]["atc"])
                        logger.info(f"✅ Added {total_atc_added} ATC documents to 'atc' section.")
                    except Exception as e:
                        logger.warning(f"⚠️ ATC AI Pass failed: {e}")

        return data

    def _preprocess_text(self, text: str) -> str:
        if not text:
            return ""
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r'Page \d+ of \d+', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\b[A-Za-z0-9]{30,}\b', '', text)
        keywords = {
            "document", "upload", "submit", "certificate", "affidavit", 
            "undertaking", "compliance", "eligibility", "experience", 
            "turnover", "performance", "annexure", "mandatory", "criteria", 
            "supporting", "proof", "report", "sheet", "form", "license", 
            "registration", "copy", "tender", "bid", "authorization",
            "declaration", "pact", "statement", "details", "record",
            "format", "acceptance", "warranty", "camc", "iso", "ca certificate",
            "net worth", "blacklisted", "statutory", "receipt", "fee", "emd",
            "security", "pbg", "bank guarantee", "letterhead", "oem", "incorporation", "gst", "pan"
        }
        filtered_lines = []
        for line in text.split('\n'):
            line = line.strip()
            if len(line) < 15 or line.startswith("---"):
                filtered_lines.append(line)
                continue
            if any(kw in line.lower() for kw in keywords):
                filtered_lines.append(line)
        text = "\n".join(filtered_lines)
        text = re.sub(r'\n{2,}', '\n', text)
        return text.strip()

    def _get_relevant_context(self, pdf_path: str) -> List[Dict]:
        scored_pages = []
        try:
            with pdfplumber.open(pdf_path) as pdf:
                total = len(pdf.pages)
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text() or ""
                    # OCR fallback omitted for brevity
                    text_lower = text.lower()
                    p_num = i + 1
                    score = 0
                    if p_num <= 5:
                        score += 100
                    keywords = [
                        "document", "upload", "submit", "certificate", "affidavit", 
                        "undertaking", "compliance", "eligibility", "experience", 
                        "turnover", "past performance", "buyer added", "annexure", 
                        "scope of work", "mandatory", "criteria", "supporting documents"
                    ]
                    for kw in keywords:
                        score += text_lower.count(kw)
                    if score > 0:
                        scored_pages.append({"p_num": p_num, "score": score, "text": f"\n--- PAGE {p_num} ---\n{text}"})

            top_pages = sorted(scored_pages, key=lambda x: x["score"], reverse=True)[:15]
            top_pages.sort(key=lambda x: x["p_num"])
            selected = [{"raw_text": p["text"], "text": self._preprocess_text(p["text"])} for p in top_pages]
            logger.info(f"📄 Filtered {total} pages down to {len(selected)}.")
        except Exception as e:
            logger.error(f"PDF Error: {e}")
        return selected

    def _fetch_atc_text(self, link: str) -> str:
        """Download ATC document with improved reliability."""
        if not link:
            return ""
        logger.info(f"🌐 Downloading ATC: {link}")
        from requests.adapters import HTTPAdapter
        from urllib3.util.retry import Retry

        session = requests.Session()
        retries = Retry(total=3, backoff_factor=1, status_forcelist=[500, 502, 503, 504])
        session.mount('http://', HTTPAdapter(max_retries=retries))
        session.mount('https://', HTTPAdapter(max_retries=retries))

        try:
            r = session.get(link, stream=True, timeout=(10, 60))
            if r.status_code != 200:
                logger.warning(f"ATC download failed with status {r.status_code}")
                return ""

            # Determine file extension
            ext = ".pdf"
            ct = r.headers.get("Content-Type", "").lower()
            cd = r.headers.get("Content-Disposition", "")

            if "filename=" in cd:
                # Extract filename from Content-Disposition
                fname_match = re.search(r'filename[^;=\n]*=(([\'"]).*?\2|[^;\n]*)', cd, re.I)
                if fname_match:
                    fname = fname_match.group(1).strip('\'"')
                    ext = os.path.splitext(fname)[1].lower()
            elif "word" in ct:
                ext = ".docx"
            elif "excel" in ct or "spreadsheet" in ct:
                ext = ".xlsx"
            elif "xml" in ct:
                ext = ".xml"

            with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
                for chunk in r.iter_content(chunk_size=8192):
                    if chunk:
                        tmp.write(chunk)
                tmp_path = tmp.name

            text_content = self._extract_text_from_file(tmp_path, ext)
            try:
                os.unlink(tmp_path)
            except Exception:
                pass
            return self._preprocess_text(text_content)

        except requests.exceptions.Timeout:
            logger.error("ATC download timed out")
            return ""
        except Exception as e:
            logger.error(f"Error fetching ATC: {e}")
            return ""

    def _extract_text_from_file(self, file_path: str, ext: str) -> str:
        text = ""
        try:
            if ext == ".pdf":
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        text += (page.extract_text() or "") + "\n"
            elif ext == ".docx":
                try:
                    from docx import Document
                    doc = Document(file_path)
                    for para in doc.paragraphs:
                        text += para.text + "\n"
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                            if row_text:
                                text += " | ".join(row_text) + "\n"
                except Exception as e:
                    logger.warning(f".docx text extraction failed: {e}")
            elif ext == ".doc":
                # .doc is old binary Word format — python-docx cannot open it.
                # Try python-docx first (works if file was mis-labelled as .doc).
                extracted = False
                try:
                    from docx import Document
                    doc = Document(file_path)
                    for para in doc.paragraphs:
                        text += para.text + "\n"
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                            if row_text:
                                text += " | ".join(row_text) + "\n"
                    extracted = True
                    logger.info(".doc opened successfully as docx")
                except Exception:
                    pass
                if not extracted:
                    # Binary fallback: extract readable ASCII text sequences from .doc
                    # Old .doc stores paragraphs as ASCII/UTF-16 runs — this finds them.
                    try:
                        with open(file_path, 'rb') as f:
                            raw = f.read()
                        import re as _re
                        # Extract all printable ASCII chunks of 5+ chars
                        chunks = _re.findall(rb'[\x20-\x7E]{5,}', raw)
                        text = ' '.join(c.decode('ascii', errors='ignore') for c in chunks)
                        logger.info(f".doc binary fallback extracted {len(text)} chars")
                    except Exception as e2:
                        logger.warning(f".doc binary extraction failed: {e2}")
            elif ext == ".xlsx":
                try:
                    import openpyxl
                    wb = openpyxl.load_workbook(file_path, data_only=True)
                    for sheet in wb:
                        for row in sheet.iter_rows(values_only=True):
                            row_str = " ".join(str(cell) for cell in row if cell)
                            if row_str:
                                text += row_str + "\n"
                except Exception as e:
                    logger.warning(f"xlsx text extraction failed: {e}")
                    pass
            elif ext == ".xml":
                import xml.etree.ElementTree as ET
                tree = ET.parse(file_path)
                for elem in tree.iter():
                    if elem.text:
                        text += elem.text + " "
            else:
                with open(file_path, 'r', errors='ignore') as f:
                    text = f.read()
        except Exception as e:
            logger.warning(f"Text extraction failed for {ext}: {e}")
        return text

    def _deep_ai_extraction(self, text: str, is_atc: bool = False) -> Dict[str, Any]:
        if not text.strip():
            return {"documents_by_section": {}}

        text = text[:30000]

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=1500,
            chunk_overlap=300,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        docs = splitter.create_documents([text])

        embedding = get_embedding_model()
        vectorstore = FAISS.from_documents(docs, embedding)
        # Increase k to 15 to grab more context for complex documents
        retriever = vectorstore.as_retriever(search_kwargs={"k": 15})

        model_selection = "llama-3.3-70b-versatile" # use 70b uniformly
        llm = ChatGroq(model=model_selection, temperature=0, groq_api_key=self.api_key)

        if is_atc:
            prompt = """You are an expert tender document reviewer. Analyze the Additional Terms and Conditions (ATC) document and identify EVERY required document, certificate, declaration, affidavit, undertaking, format, annexure, or attachment the bidder needs to submit.

Pay SPECIAL attention to:
- "Annexure" forms (e.g., Annexure 1, 2, 3, etc.)
- "Format" checklists (e.g., Format 1, Format 2)
- Certificates (ISO, CA Certificate for Turnover/Net Worth, Company Registration, GST, PAN)
- Acceptance letters (e.g., Acceptance of ATC on letterhead)
- Undertakings & Affidavits (e.g., Not Blacklisted, Statutory, Warranty/CAMC)
- Receipts (e.g., EMD payment receipt)
- OEM Compliance and Authorization

Return ONLY a JSON object:
{"documents_by_section": {"atc": ["Document Name 1", "Document Name 2", ...]}}
"""
        else:
            prompt = """You are an expert tender document reviewer. Extract ALL required bidder documents from the tender.

Pay SPECIAL attention to required Annexures, Formats, Checklists, Acceptance of ATC, ISO, CA Certificates, Affidavits, and Specific Undertakings.

Categorize into:
- main (from 'Document required from seller')
- financial
- technical
- certificates
- additional

Return ONLY a JSON object returning arrays of strings:
{"documents_by_section": {"main": [], "financial": [], "technical": [], "certificates": [], "additional": []}}
"""

        qa = RetrievalQA.from_chain_type(llm=llm, retriever=retriever)

        response = ""
        for attempt in range(3):
            try:
                response = qa.invoke({"query": prompt})["result"]
                break
            except Exception as e:
                logger.warning(f"Retry {attempt+1}: {e}")
                time.sleep(2 ** attempt)

        if not response:
            logger.warning(f"No response from LLM for is_atc={is_atc}")
            return {"documents_by_section": {}}

        logger.info(f"LLM Response (is_atc={is_atc}):\n{response[:500]}...\n")
        parsed = self._extract_json_from_response(response)
        if not parsed:
            logger.warning(f"Failed to parse JSON from LLM response (is_atc={is_atc})")
            return {"documents_by_section": {}}

        cleaned = {}
        for section, docs in parsed.get("documents_by_section", {}).items():
            if not isinstance(docs, list):
                continue
            cleaned[section] = []
            seen = set()
            for d in docs:
                if not isinstance(d, str):
                    continue
                d = re.sub(r'^[\s\d\.\*\-–—]+', '', d).strip()
                d = re.sub(r'\s+', ' ', d)
                if 3 <= len(d) <= 100:
                    key = d.lower().translate(str.maketrans('', '', '()[]{}.,- '))
                    if key and key not in seen:
                        seen.add(key)
                        cleaned[section].append(d)

        if is_atc and "atc" not in cleaned:
            all_docs = []
            for docs in cleaned.values():
                all_docs.extend(docs)
            if all_docs:
                cleaned = {"atc": all_docs}

        logger.info(f"AI extraction: {sum(len(v) for v in cleaned.values())} documents")
        return {"documents_by_section": cleaned}

    def _extract_json_from_response(self, response: str) -> Optional[Dict]:
        if not response:
            return None
        response = response.strip()
        strategies = [
            lambda r: re.search(r'```json\s*(\{.*?\})\s*```', r, re.DOTALL),
            lambda r: re.search(r'```\s*(\{.*?\})\s*```', r, re.DOTALL),
            lambda r: re.search(r'(\{.*\})', r, re.DOTALL),
        ]
        json_str = None
        for strategy in strategies:
            match = strategy(response)
            if match:
                json_str = match.group(1).strip()
                break
        if not json_str:
            start = response.find('{')
            if start == -1:
                return None
            balance = 0
            end = start
            for i, ch in enumerate(response[start:], start):
                if ch == '{':
                    balance += 1
                elif ch == '}':
                    balance -= 1
                    if balance == 0:
                        end = i
                        break
            if balance == 0:
                json_str = response[start:end+1]
            else:
                return None
        try:
            return json.loads(json_str)
        except json.JSONDecodeError:
            json_str = re.sub(r',\s*([}\]])', r'\1', json_str)
            json_str = re.sub(r'([{,])\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*:', r'\1"\2":', json_str)
            try:
                return json.loads(json_str)
            except Exception:
                return None


    # ============================================================
    # REGEX FALLBACK
    # ============================================================

    def _regex_fallback(self, text: str, data: Dict) -> Dict:
        """Pure regex extraction when AI is unavailable or fails"""

        sections = {
            "main": [], "atc": [], "financial": [],
            "technical": [], "certificates": [], "additional": []
        }

        # MAIN: "Document required from seller" field
        doc_match = re.search(
            r'Document required from seller[:\s/]*\n?(.*?)(?=\*In case|\n\n|\Z)',
            text, re.IGNORECASE | re.DOTALL
        )
        if doc_match:
            for item in re.split(r'[,\n]+', doc_match.group(1)):
                item = item.strip().strip('*').strip()
                if item and len(item) > 2 and len(item) < 150:
                    sections["main"].append(item)

        # ATC: look for numbered items in buyer-added ATC
        atc_match = re.search(
            r'Buyer Added Bid Specific.*?(.*?)(?=Disclaimer|अस्वीकरण|\Z)',
            text, re.IGNORECASE | re.DOTALL
        )
        if atc_match:
            numbered = re.findall(r'\d+\.\s+([A-Z][^.\n]{10,100})', atc_match.group(1))
            for item in numbered:
                if any(kw in item.lower() for kw in [
                    'upload', 'submit', 'certificate', 'affidavit', 'undertaking',
                    'declaration', 'document', 'sheet', 'proof', 'pact', 'form',
                    'brochure', 'data sheet', 'letter', 'copy', 'mandate'
                ]):
                    sections["atc"].append(item.strip()[:100])

        # FINANCIAL
        financial_patterns = [
            (r'EMD Amount', "EMD (Earnest Money Deposit)"),
            (r'Performance.*?Bank Guarantee|ePBG', "Performance Bank Guarantee (PBG)"),
            (r'Audited Balance Sheet', "Audited Balance Sheets"),
            (r'CA.*?turnover|Chartered Accountant.*?turnover', "CA Certificate for Turnover"),
            (r'Profit\s*[&]\s*Loss|Profit and Loss', "Profit & Loss Statement"),
            (r'Income Tax Acknowledgement', "Income Tax Acknowledgement"),
            (r'Net Worth', "Net Worth Certificate"),
            (r'ITR.*?\d{4}', "Income Tax Returns"),
        ]
        for pat, name in financial_patterns:
            if re.search(pat, text, re.IGNORECASE):
                sections["financial"].append(name)

        # TECHNICAL
        technical_patterns = [
            (r'Data Sheet.*?product|product.*?data sheet', "Product Data Sheet"),
            (r'catalogue|brochure', "Product Catalogue / Brochure"),
            (r'Technical.*?Compliance.*?Sheet', "Technical Compliance Sheet"),
            (r'Compliance.*?BoQ', "BoQ Compliance Document"),
            (r'Technical Bid', "Technical Bid Documents"),
        ]
        for pat, name in technical_patterns:
            if re.search(pat, text, re.IGNORECASE):
                sections["technical"].append(name)

        # CERTIFICATES
        certificate_patterns = [
            (r'ISO\s*(?:9001)?', "ISO Certificate"),
            (r'\bBIS\b', "BIS Certificate"),
            (r'Drug License', "Drug License"),
            (r'Udyam Registration', "Udyam Registration Certificate"),
            (r'DPIIT.*?Startup|Startup.*?DPIIT', "DPIIT Startup Certificate"),
            (r'GST.*?Registration|GSTIN|Copy of GST', "GST Registration Certificate"),
            (r'PAN Card|Copy of PAN', "PAN Card"),
            (r'MSME', "MSME/Udyam Certificate"),
            (r'Certificate of (?:Company )?Registration|Certificate of Incorporation', "Company Registration Certificate"),
        ]
        for pat, name in certificate_patterns:
            if re.search(pat, text, re.IGNORECASE):
                sections["certificates"].append(name)

        # ADDITIONAL
        additional_patterns = [
            (r'Land Border Declaration', "Land Border Declaration"),
            (r'Integrity Pact', "Integrity Pact"),
            (r'MII.*?certif|Make.*?India.*?certif', "MII Local Content Certificate"),
            (r'Malicious Code', "Malicious Code Certificate (MCC)"),
            (r'EFT Mandate', "EFT Mandate"),
            (r'Statutory Undertaking', "Statutory Undertaking"),
            (r'Warranty.*?CAMC|CAMC.*?Warranty', "Undertaking for Warranty/CAMC"),
            (r'Acceptance of ATC', "Acceptance of ATC"),
            (r'Format\s*1', "Format 1 - Checklist"),
            (r'Format\s*2', "Format 2 - Particulars of Bidding Agency"),
            (r'Annexure\s*[1I]\b', "Annexure 1"),
            (r'Annexure\s*[2II]\b', "Annexure 2 - Compliance to MTS"),
            (r'Annexure\s*[3III]\b', "Annexure 3 - Statutory Undertaking"),
            (r'Annexure\s*[4IV]\b', "Annexure 4 - Undertaking for Warranty/CAMC"),
            (r'Annexure\s*XI\b', "Annexure XI - Undertaking for Relaxations"),
            (r'Undertaking', "Undertaking Letter"),
            (r'Affidavit', "Affidavit"),
        ]
        for pat, name in additional_patterns:
            if re.search(pat, text, re.IGNORECASE):
                sections["additional"].append(name)

        # Deduplicate within sections
        for sec in sections:
            sections[sec] = self._deduplicate(sections[sec])

        data["documents_by_section"] = sections
        return data

    # ============================================================
    # BASIC INFO (regex, fast)
    # ============================================================

    def _extract_basic_info(self, text: str) -> Dict[str, Any]:
        data = {
            "bid_number": None,
            "bid_end_date": None,
            "item_category": None,
            "classification_level": "UNIVERSAL",
            "turnover": 0,
            "oem_turnover": 0,
            "experience": 0,
            "past_performance": 0,
            "mse_status": "No",
            "startup_status": "No",
            "emd_amount": None,
            "documents_by_section": {
                "main": [], "atc": [], "financial": [],
                "technical": [], "certificates": [], "additional": []
            },
            "keyword_checks": {},
            "atc_document_link": None
        }

        # Bid Number
        m = re.search(r'GEM/\d{4}/[A-Z]/\d+', text)
        if m:
            data["bid_number"] = m.group()

        # Bid End Date
        m = re.search(r'Bid End Date[/\s:]*(?:Time)?\s*(\d{2}-\d{2}-\d{4})', text, re.IGNORECASE)
        if m:
            data["bid_end_date"] = m.group(1)

        # Item Category
        m = re.search(r'Item Category\s*[:\n]+\s*([^\n]+)', text, re.IGNORECASE)
        if m:
            raw_cat = m.group(1).strip()
            q = re.search(r'\(Q([1-4])\)', raw_cat)
            if q:
                data["classification_level"] = f"Q{q.group(1)}"
                raw_cat = re.sub(r'\s*\(Q[1-4]\)', '', raw_cat).strip()
            data["item_category"] = raw_cat

        if data["classification_level"] == "UNIVERSAL":
            # Robust multi-pattern Q1-Q4 fallback detection
            # Pattern A: labeled keyword followed by Q1-Q4 on the same line
            cl_match = re.search(
                r'(?:Item\s+Categ|Bid\s+Categ|Quadrant|Classification|Bid\s+Type|Bidding\s+Type)'
                r'[^\n]{0,80}?\b(Q[1-4])\b',
                text, re.IGNORECASE
            )

            # Pattern B: Q1-Q4 inside parentheses — very common: "(Q1)"
            if not cl_match:
                cl_match = re.search(r'\((Q[1-4])\)', text, re.IGNORECASE)

            # Pattern C: line that STARTS with Q1/Q2/Q3/Q4 (short label lines)
            if not cl_match:
                for line in text.splitlines():
                    ls = line.strip()
                    m2 = re.match(r'^(Q[1-4])\b', ls, re.IGNORECASE)
                    if m2 and len(ls) < 10:
                        cl_match = m2
                        break

            if cl_match:
                data["classification_level"] = cl_match.group(1).upper()
                logger.info(f"[CLASSIFY] Fallback regex matched: {data['classification_level']}")

        # Turnover
        m = re.search(
            r'Minimum Average Annual Turnover.*?(\d+(?:\.\d+)?)\s*(Lakh|Crore)',
            text, re.IGNORECASE
        )
        if m:
            val, unit = float(m.group(1)), m.group(2).lower()
            data["turnover"] = int(val * (10_000_000 if unit == "crore" else 100_000))

        # Experience
        m = re.search(r'Years of Past Experience Required.*?(\d+)\s*Year', text, re.IGNORECASE)
        if m:
            data["experience"] = int(m.group(1))

        # Past Performance
        m = re.search(r'Past Performance[:\s/]*(\d+)\s*%', text, re.IGNORECASE)
        if m:
            data["past_performance"] = int(m.group(1))

        # MSE Status
        m = re.search(r'MSE (?:Exemption|Relaxation) for Years.*?(Yes|No)', text, re.IGNORECASE)
        if m:
            data["mse_status"] = "Yes" if m.group(1).lower() == "yes" else "No"

        # Startup Status
        m = re.search(r'Startup (?:Exemption|Relaxation) for Years.*?(Yes|No)', text, re.IGNORECASE)
        if m:
            data["startup_status"] = "Yes" if m.group(1).lower() == "yes" else "No"

        # EMD Amount
        m = re.search(r'EMD Amount[:\s]*(\d+)', text, re.IGNORECASE)
        if m:
            data["emd_amount"] = int(m.group(1))

        return data

    # ============================================================
    # KEYWORD CHECKS
    # ============================================================

    def _check_keywords(self, text: str) -> Dict[str, bool]:
        results = {}
        for keyword, patterns in SPECIAL_KEYWORDS.items():
            found = any(re.search(p, text, re.IGNORECASE) for p in patterns)
            results[keyword] = found
            logger.info(f"  {'✅' if found else '❌'} {keyword}")
        return results

    # ============================================================
    # ATC LINK
    # ============================================================

    def _extract_atc_link(self, pdf_path: str) -> Optional[str]:
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text() or ""
                    if re.search(r"Buyer uploaded ATC document|Click here to view", text, re.IGNORECASE):
                        if page.hyperlinks:
                            for hl in page.hyperlinks:
                                uri = hl.get("uri", "")
                                if "http" in uri:
                                    logger.info(f"🔗 ATC link: {uri}")
                                    return uri
        except Exception as e:
            logger.error(f"ATC link error: {e}")
        return None

    # ============================================================
    # UTILITIES
    # ============================================================

    def _deduplicate(self, docs: List[str]) -> List[str]:
        seen, result = set(), []
        for doc in docs:
            doc = doc.strip()
            if not doc or len(doc) < 3:
                continue
            key = re.sub(r'[^a-z0-9]', '', doc.lower())
            if key and len(key) > 3 and key not in seen:
                seen.add(key)
                result.append(doc)
        return result

    def _default_data(self) -> Dict[str, Any]:
        return {
            "bid_number": None,
            "bid_end_date": None,
            "item_category": None,
            "classification_level": "UNIVERSAL",
            "turnover": 0,
            "oem_turnover": 0,
            "experience": 0,
            "past_performance": 0,
            "mse_status": "No",
            "startup_status": "No",
            "emd_amount": None,
            "documents_by_section": {
                "main": [], "atc": [], "financial": [],
                "technical": [], "certificates": [], "additional": []
            },
            "keyword_checks": {k: False for k in SPECIAL_KEYWORDS},
            "atc_document_link": None
        }


# ============================================================
# CLI TEST
# ============================================================

if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python extractor.py <path_to_pdf>")
        sys.exit(1)

    extractor = CompleteTenderExtractor()
    result = extractor.extract_fields(sys.argv[1])

    print("\n" + "=" * 80)
    print("📄 GeM TENDER EXTRACTION RESULT")
    print("=" * 80)
    
    # Print basic fields
    basic_fields = ["bid_number", "bid_end_date", "item_category", "classification_level",
                    "turnover", "experience", "past_performance", "mse_status",
                    "startup_status", "emd_amount"]
    
    for field in basic_fields:
        val = result.get(field, "N/A")
        if field == "turnover" and val:
            val = f"₹{val:,}"
        print(f"  {field:<22}: {val}")

    print(f"\n  ATC Link: {result.get('atc_document_link', 'None')}")

    # Print documents by section
    print("\n📂 DOCUMENTS BY SECTION:")
    labels = {
        "main":         "📋 Main (Document required from seller)",
        "atc":          "📄 ATC / Buyer Added Terms",
        "financial":    "💰 Financial",
        "technical":    "🔧 Technical",
        "certificates": "📜 Certificates",
        "additional":   "📎 Additional",
    }
    
    for sec, docs in result.get("documents_by_section", {}).items():
        if not docs: continue
        label = labels.get(sec, f"📌 {sec.replace('_', ' ').title()}")
        print(f"\n  {label} ({len(docs)}):")
        for d in docs:
            print(f"     • {d}")

    # Flatten and calculate total unique documents
    unique_docs = set()
    for sec_docs in result.get("documents_by_section", {}).values():
        for doc in sec_docs:
            unique_docs.add(doc.lower().strip())
    
    print(f"\n📊 TOTAL UNIQUE DOCUMENTS (across all sections): {len(unique_docs)}")

    # Print keyword checks
    print("\n🔍 KEYWORD CHECKS:")
    for keyword, found in result.get("keyword_checks", {}).items():
        print(f"  {'✅' if found else '❌'} {keyword}")

    # Save to JSON file
    with open("extraction_result.json", "w", encoding="utf-8") as f:
        json.dump(result, f, indent=2, ensure_ascii=False, default=str)
    print("\n✅ Saved to extraction_result.json")